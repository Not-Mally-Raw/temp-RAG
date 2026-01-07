"""High-performance manufacturing rule extraction pipeline.

This module consolidates document loading, chunking, LLM interaction, and
post-processing so the rest of the codebase can stay lean.  It replaces the
previous mix of *enhanced* modules while remaining compatible with the
``FastRuleExtractor``/``ProductionRuleExtractionSystem`` interfaces.
"""

from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import time
from collections import OrderedDict
from dataclasses import dataclass, field, replace
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import tiktoken
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, SystemMessage
from pydantic import Field
from pydantic_settings import BaseSettings

from .prompts import PromptContext, PromptLibrary

try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover - library optional
    fitz = None

try:  # DOCX extraction
    import docx  # type: ignore
except Exception:  # pragma: no cover - library optional
    docx = None

try:  # DataFrame powered formats
    import pandas as pd  # type: ignore
except Exception:  # pragma: no cover - library optional
    pd = None

try:  # JSON readability metrics
    import textstat  # type: ignore
except Exception:  # pragma: no cover - library optional
    textstat = None

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Configuration dataclasses
# ---------------------------------------------------------------------------


class RuleExtractionSettings(BaseSettings):
    """Global settings used across the fast extraction pipeline."""

    groq_api_key: str = Field(default="")
    groq_model: str = Field(default="meta-llama/llama-4-scout-17b-16e-instruct")
    temperature: float = Field(default=0.1)
    max_output_tokens: int = Field(default=2048)

    max_chunk_tokens: int = Field(default=3200)
    chunk_overlap_tokens: int = Field(default=200)
    max_rules_per_chunk: int = Field(default=20)
    max_rules_total: int = Field(default=150)

    max_concurrent_calls: int = Field(default=4)
    throttle_seconds: float = Field(default=0.0)
    request_timeout: float = Field(default=60.0)
    max_retries: int = Field(default=3)
    retry_backoff_seconds: float = Field(default=2.0)

    max_documents_concurrency: int = Field(default=3)
    min_text_length: int = Field(default=64)
    document_cache_size: int = Field(default=16)

    class Config:
        env_file = ".env"
        extra = "ignore"


# ---------------------------------------------------------------------------
# Lightweight data containers
# ---------------------------------------------------------------------------


@dataclass(slots=True)
class Rule:
    """Synthetic representation of a manufacturing rule."""

    rule_text: str
    category: str
    rule_type: str
    confidence: float
    priority: str = "medium"
    rationale: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        payload = {
            "rule_text": self.rule_text.strip(),
            "rule_category": self.category,
            "rule_type": self.rule_type,
            "confidence_score": round(float(self.confidence), 4),
            "priority": self.priority,
            "rationale": self.rationale,
        }
        payload.update(self.metadata)
        return payload


@dataclass(slots=True)
class DocumentPayload:
    """Text plus metadata returned by :class:`DocumentLoader`."""

    text: str
    metadata: Dict[str, Any]


@dataclass(slots=True)
class ExtractionSummary:
    """Result payload shared with existing public APIs."""

    filename: str
    status: str
    rules: List[Dict[str, Any]]
    rule_count: int
    chunks_processed: int
    model_calls: int
    processing_time: float
    avg_confidence: float
    document_metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "filename": self.filename,
            "status": self.status,
            "rules": self.rules,
            "rule_count": self.rule_count,
            "chunks_processed": self.chunks_processed,
            "model_calls": self.model_calls,
            "processing_time": self.processing_time,
            "avg_confidence": self.avg_confidence,
            "document_metadata": self.document_metadata,
        }


# ---------------------------------------------------------------------------
# Support classes – rate limiting, chunking, document loading
# ---------------------------------------------------------------------------


class AsyncRateLimiter:
    """Coarse global rate limiter to avoid hammering the Groq API."""

    def __init__(self, min_interval: float) -> None:
        self._interval = max(0.0, float(min_interval))
        self._lock = asyncio.Lock()
        self._last_call = 0.0

    async def wait(self) -> None:
        if self._interval <= 0:
            return

        async with self._lock:
            now = time.perf_counter()
            delta = now - self._last_call
            if delta < self._interval:
                await asyncio.sleep(self._interval - delta)
            self._last_call = time.perf_counter()

    def set_interval(self, interval: float) -> None:
        self._interval = max(0.0, float(interval))

    def get_interval(self) -> float:
        return self._interval


class TextChunker:
    """Token-aware chunker with quick word-based fallback."""

    def __init__(self, encoding_name: str = "cl100k_base") -> None:
        self._encoding_name = encoding_name
        self._encoding = tiktoken.get_encoding(encoding_name)

    def split(self, text: str, *, window: int, overlap: int) -> List[str]:
        if not text:
            return []

        tokens = self._encoding.encode(text)
        if not tokens:
            return []

        window = max(1, int(window))
        overlap = max(0, min(int(overlap), window - 1))
        step = window - overlap

        chunks: List[str] = []
        for start in range(0, len(tokens), step):
            end = min(start + window, len(tokens))
            chunk_tokens = tokens[start:end]
            chunks.append(self._encoding.decode(chunk_tokens))
            if end == len(tokens):
                break
        return chunks or [text]

    def word_fallback(self, text: str, *, window: int, overlap: int) -> List[str]:
        words = text.split()
        if not words:
            return []
        limit = max(1, int(window / 1.3))
        stride = max(1, limit - int(overlap / 1.3))
        return [" ".join(words[i : i + limit]) for i in range(0, len(words), stride)]


class DocumentLoader:
    """Cached multi-format document loader."""

    def __init__(self, cache_size: int = 16) -> None:
        self._cache_size = max(1, cache_size)
        self._cache: "OrderedDict[str, DocumentPayload]" = OrderedDict()
        self._cache_lock = asyncio.Lock()

    async def load(self, file_path: Path) -> DocumentPayload:
        file_path = file_path.expanduser().resolve()
        cache_key = await self._cache_key(file_path)

        async with self._cache_lock:
            cached = self._cache.get(cache_key)
            if cached:
                self._cache.move_to_end(cache_key)
                return cached

        payload = await asyncio.to_thread(self._load_sync, file_path)

        async with self._cache_lock:
            self._cache[cache_key] = payload
            if len(self._cache) > self._cache_size:
                self._cache.popitem(last=False)
        return payload

    async def _cache_key(self, file_path: Path) -> str:
        stat = await asyncio.to_thread(file_path.stat)
        return f"{file_path}|{stat.st_mtime_ns}|{stat.st_size}"

    # ------------------------------------------------------------------
    # Sync helpers run inside ``asyncio.to_thread``
    # ------------------------------------------------------------------

    def _load_sync(self, file_path: Path) -> DocumentPayload:
        suffix = file_path.suffix.lower()
        loader = {
            ".pdf": self._load_pdf,
            ".txt": self._load_text,
            ".md": self._load_text,
            ".json": self._load_json,
            ".docx": self._load_docx,
            ".csv": self._load_table,
            ".xlsx": self._load_table,
            ".xls": self._load_table,
        }.get(suffix, self._load_text)

        text, extra = loader(file_path)
        metadata = {
            "filename": file_path.name,
            "file_path": str(file_path),
            "file_extension": suffix,
            "char_count": len(text),
            "word_count": len(text.split()),
            "page_count": extra.get("page_count"),
            "processing_method": extra.get("processing_method"),
            "readability_score": None,
            "complexity_score": None,
            "manufacturing_keywords": self._detect_keywords(text),
        }

        if textstat is not None and text:
            try:
                metadata["readability_score"] = textstat.flesch_reading_ease(text)
                metadata["complexity_score"] = textstat.flesch_kincaid_grade(text)
            except Exception:  # pragma: no cover - textstat edge cases
                pass

        metadata.update(extra)
        return DocumentPayload(text=text, metadata=metadata)

    # ---------------------------- format loaders -----------------------------

    def _load_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if fitz is None:
            raise RuntimeError("PyMuPDF (fitz) is required for PDF extraction")

        doc = fitz.open(file_path)
        text_parts: List[str] = []
        remove_hyphen_flag = getattr(
            fitz,
            "TEXT_DEHYPHENATE",
            getattr(fitz, "TEXTFLAGS_REMOVE_HYPHENS", 0),
        )
        # Prefer the modern TEXT_DEHYPHENATE constant but keep backward compatibility.
        for index in range(len(doc)):
            page = doc.load_page(index)
            text = page.get_text("text", flags=remove_hyphen_flag)
            if text.strip():
                text_parts.append(f"--- Page {index + 1} ---\n{text}")
        doc.close()
        return "\n\n".join(text_parts), {
            "processing_method": "pymupdf",
            "page_count": len(text_parts),
        }

    def _load_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if docx is None:
            raise RuntimeError("python-docx is required for DOCX extraction")

        document = docx.Document(str(file_path))
        parts: List[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_values:
                    parts.append(" | ".join(row_values))
        return "\n\n".join(parts), {"processing_method": "python-docx"}

    def _load_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        try:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = file_path.read_text(errors="ignore")
        return text, {"processing_method": "text"}

    def _load_json(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        raw = json.loads(file_path.read_text(encoding="utf-8"))
        return json.dumps(raw, indent=2, ensure_ascii=False), {"processing_method": "json"}

    def _load_table(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        if pd is None:
            raise RuntimeError("pandas is required for CSV/Excel extraction")

        text_parts: List[str] = []
        if file_path.suffix.lower() == ".csv":
            df = pd.read_csv(file_path)
            frames = {"CSV": df}
        else:
            excel = pd.ExcelFile(file_path)
            frames = {name: excel.parse(name) for name in excel.sheet_names}

        for sheet, frame in frames.items():
            if frame.empty:
                continue
            preview = frame.head(100)
            text_parts.append(f"=== Sheet: {sheet} ===")
            text_parts.append(" | ".join(map(str, preview.columns)))
            for _, row in preview.iterrows():
                values = [str(value) if pd.notna(value) else "" for value in row.tolist()]
                text_parts.append(" | ".join(values))

        return "\n".join(text_parts), {"processing_method": "pandas"}

    # --------------------------- helper utilities ---------------------------

    def _detect_keywords(self, text: str) -> List[str]:
        if not text:
            return []
        glossary = {
            "steel",
            "aluminum",
            "plastic",
            "composite",
            "stainless steel",
            "machining",
            "welding",
            "casting",
            "forging",
            "stamping",
            "tolerance",
            "specification",
            "inspection",
            "assembly",
            "fastener",
            "fixture",
            "cycle time",
            "surface finish",
        }
        lower = text.lower()
        return sorted({term for term in glossary if term in lower})


# ---------------------------------------------------------------------------
# LLM interaction and post-processing
# ---------------------------------------------------------------------------


class ChunkProcessor:
    """Handles prompt rendering, API calls, and response parsing."""

    def __init__(self, settings: RuleExtractionSettings) -> None:
        if not settings.groq_api_key:
            settings.groq_api_key = os.getenv("GROQ_API_KEY", "")
        if not settings.groq_api_key:
            raise ValueError("GROQ_API_KEY missing; set it in the environment or .env")

        self.settings = settings
        self.prompts = PromptLibrary()
        # Optional mock mode for local/offline development: set ALLOW_FAKE_GROQ=1
        allow_fake = os.getenv("ALLOW_FAKE_GROQ", "0") == "1"
        self._mock_mode = allow_fake
        if allow_fake:
            import types, json as _json

            async def _fake_ainvoke(messages):  # type: ignore
                # Extract last human message to synthesise simple rules
                human_content = ""
                for m in reversed(messages):
                    if hasattr(m, "content"):
                        human_content = getattr(m, "content", "")
                        break
                snippet = human_content[:200].replace("\n", " ")
                fake_payload = {
                    "rules": [
                        {
                            "rule_text": "Calibrate equipment before operation",
                            "rule_category": "Setup",
                            "rule_type": "Mandatory",
                            "confidence": 0.95,
                            "priority": "high",
                            "rationale": "Baseline mock rule",
                            "primary_feature": "calibration",
                            "supporting_quote": snippet,
                            "supporting_quotes": [snippet],
                            "intent": "Calibration readiness",
                            "scope_domain": {"process": "Assembly", "item": "Equipment", "feature": None},
                            "applicability": [],
                            "constraints": [
                                {
                                    "subject": "calibration_status",
                                    "operator": "==",
                                    "value": "calibrated",
                                    "unit": None,
                                    "logic": "qualitative",
                                    "provenance_quote": snippet,
                                }
                            ],
                            "severity": "ENFORCEABLE",
                            "validation_state": "ENFORCEABLE",
                        },
                        {
                            "rule_text": "Ensure safety guards engaged",
                            "rule_category": "Safety",
                            "rule_type": "Mandatory",
                            "confidence": 0.9,
                            "priority": "high",
                            "rationale": "Baseline mock rule",
                            "primary_feature": "safety",
                            "supporting_quote": snippet,
                            "supporting_quotes": [snippet],
                            "intent": "Safety guarding",
                            "scope_domain": {"process": "General", "item": "Equipment", "feature": "Guard"},
                            "applicability": [],
                            "constraints": [
                                {
                                    "subject": "safety_guard_state",
                                    "operator": "==",
                                    "value": "engaged",
                                    "unit": None,
                                    "logic": "qualitative",
                                    "provenance_quote": snippet,
                                }
                            ],
                            "severity": "ENFORCEABLE",
                            "validation_state": "ENFORCEABLE",
                        },
                    ],
                }
                return types.SimpleNamespace(content=f"```json\n{_json.dumps(fake_payload)}\n```")

            # Minimal client shim with ainvoke attribute
            self._client = types.SimpleNamespace(ainvoke=_fake_ainvoke)
        else:
            self._client = ChatGroq(
                model=settings.groq_model,
                api_key=settings.groq_api_key,
                temperature=settings.temperature,
                max_tokens=settings.max_output_tokens,
                timeout=settings.request_timeout,
            )
        self._rate_limiter = AsyncRateLimiter(settings.throttle_seconds)
        self.system_prompt = self.prompts.system_prompt

    async def process(self, chunk_text: str, *, document_name: str, chunk_index: int) -> List[Rule]:
        await self._rate_limiter.wait()
        prompt = self.prompts.build_fast_prompt(
            PromptContext(
                document_name=document_name,
                chunk_index=chunk_index,
                chunk_text=chunk_text,
            )
        )

        backoff = self.settings.retry_backoff_seconds
        last_error_message = ""
        for attempt in range(1, self.settings.max_retries + 1):
            try:
                if self._mock_mode:
                    # In mock mode we bypass network and just call shim directly (no wait_for timeout enforcement)
                    response = await self._client.ainvoke([
                        SystemMessage(content=self.system_prompt),
                        HumanMessage(content=prompt),
                    ])
                else:
                    response = await asyncio.wait_for(
                        self._client.ainvoke(
                            [
                                SystemMessage(content=self.system_prompt),
                                HumanMessage(content=prompt),
                            ]
                        ),
                        timeout=self.settings.request_timeout,
                    )
                raw = response.content if hasattr(response, "content") else str(response)
                return self._parse_response(raw, document_name, chunk_index)
            except asyncio.TimeoutError:
                logger.warning(
                    "model_timeout",
                    extra={"chunk_index": chunk_index, "attempt": attempt},
                )
                last_error_message = "timeout"
            except Exception as error:  # pragma: no cover - network edge cases
                last_error_message = str(error)
                if "429" in last_error_message:
                    retry_after = None
                    response = getattr(error, "response", None)
                    if response is not None:
                        headers = getattr(response, "headers", {})
                        retry_after = headers.get("retry-after") if headers else None
                    new_interval = max(
                        self._rate_limiter.get_interval(),
                        self.settings.throttle_seconds or 0.0,
                        1.5,
                    )
                    if retry_after:
                        try:
                            new_interval = max(new_interval, float(retry_after))
                        except (TypeError, ValueError):
                            pass
                    new_interval = max(new_interval, 2.0)
                    self._rate_limiter.set_interval(new_interval)
                    self.settings.throttle_seconds = new_interval
                # Avoid using reserved LogRecord attribute name 'message'
                logger.warning(
                    "model_call_failed",
                    extra={
                        "chunk_index": chunk_index,
                        "attempt": attempt,
                        "error_message": last_error_message,
                    },
                )
            if attempt == self.settings.max_retries:
                break
            delay = backoff * attempt
            if "429" in last_error_message:
                delay = max(delay, self.settings.throttle_seconds or 1.5)
            await asyncio.sleep(delay)
        return []

    # ------------------------- JSON parsing helpers -------------------------

    def _parse_response(self, raw: str, document_name: str, chunk_index: int) -> List[Rule]:
        json_blob = self._extract_json(raw)
        if not json_blob:
            return []

        try:
            payload = json.loads(json_blob)
        except json.JSONDecodeError:
            logger.debug("json_decode_failed", extra={"chunk_index": chunk_index})
            return []

        items: Iterable[Any]
        if isinstance(payload, dict) and "rules" in payload:
            items = payload["rules"]
        elif isinstance(payload, list):
            items = payload
        else:
            return []

        clean: List[Rule] = []
        for idx, item in enumerate(items):
            rule = self._normalise_rule(item, document_name, chunk_index, idx)
            if rule:
                clean.append(rule)
            if len(clean) >= self.settings.max_rules_per_chunk:
                break
        return clean

    @staticmethod
    def _render_prompt(chunk_text: str, document_name: str, chunk_index: int) -> str:
        return (
            f"Document: {document_name}\n"
            f"Segment: {chunk_index + 1}\n"
            "Instructions: Extract high-value manufacturing rules as a JSON object with a key named rules. "
            "The value must be a list of rule objects. Each rule object must contain rule_text, rule_category, "
            "rule_type, confidence (0-1 float), priority (high|medium|low), rationale, primary_feature, unit, "
            "value, tolerance_range, supporting_quote. Only return JSON.\n"
            "Source text:\n"
            f"{chunk_text}"
        )

    @staticmethod
    def _extract_json(response_text: str) -> str:
        text = response_text.strip()
        if text.startswith("```"):
            match = re.search(r"```json\s*(.*?)```", text, re.DOTALL)
            if match:
                return match.group(1).strip()
            match = re.search(r"```\s*(.*?)```", text, re.DOTALL)
            if match:
                return match.group(1).strip()

        text_no_bom = text.replace("\ufeff", "").strip()
        if text_no_bom.startswith("[") and text_no_bom.endswith("]"):
            return text_no_bom
        if text_no_bom.startswith("{") and text_no_bom.endswith("}"):
            return text_no_bom

        match = re.search(r"(\{.*\}|\[.*\])", text_no_bom, re.DOTALL)
        if match:
            return match.group(1).strip()
        return ""

    def _normalise_rule(
        self,
        payload: Any,
        document_name: str,
        chunk_index: int,
        rule_index: int,
    ) -> Optional[Rule]:
        if not isinstance(payload, dict):
            return None

        canonical = str(payload.get("canonical_statement", "")).strip()
        raw_text = str(payload.get("rule_text", "")).strip()
        if not raw_text and canonical:
            raw_text = canonical
        if not raw_text:
            return None

        category = str(payload.get("rule_category", "General")).strip() or "General"
        rule_type = str(payload.get("rule_type", "general")).strip() or "general"
        priority = str(payload.get("priority", "medium")).lower()
        if priority not in {"low", "medium", "high"}:
            priority = "medium"

        confidence_raw = (
            payload.get("confidence")
            or payload.get("confidence_score")
            or payload.get("confidenceScore")
            or 0.6
        )
        try:
            confidence = float(confidence_raw)
        except (TypeError, ValueError):
            confidence = 0.6
        confidence = max(0.0, min(1.0, confidence))

        intent = str(
            payload.get("intent")
            or payload.get("rule_intent")
            or payload.get("ruleIntent")
            or ""
        ).strip()

        scope_domain = payload.get("scope_domain") or payload.get("scope") or payload.get("domain")
        if not isinstance(scope_domain, dict):
            scope_domain = {"process": None, "item": None, "feature": None}
        else:
            scope_domain = {
                "process": scope_domain.get("process"),
                "item": scope_domain.get("item"),
                "feature": scope_domain.get("feature"),
            }

        # Preserve contract-style domain strings (e.g. "sheet metal / Bend") separately.
        domain_text = payload.get("domain")
        if not isinstance(domain_text, str):
            domain_text = ""

        # Applicability: prefer structured object; fall back to list/conditions.
        app_struct = payload.get("applicability")
        applicability: List[Any]
        if isinstance(app_struct, dict):
            applicability = [app_struct]
        else:
            applicability = (
                payload.get("applicability")
                or payload.get("applicability_conditions")
                or payload.get("conditions")
                or []
            )
            if not isinstance(applicability, list):
                applicability = []

        # Constraints: Level-1 may emit a string expression or a structured list.
        constraints_field = payload.get("constraints") or payload.get("requirements")
        constraints_is_string = isinstance(constraints_field, str)
        constraints_raw = constraints_field if constraints_is_string else (constraints_field or [])
        if not constraints_is_string and not isinstance(constraints_raw, list):
            constraints_raw = []

        primary_constraints = payload.get("primary_constraints") or []
        harvested_constraints = payload.get("harvested_constraints") or []
        if not isinstance(primary_constraints, list):
            primary_constraints = []
        if not isinstance(harvested_constraints, list):
            harvested_constraints = []

        merged_constraints: List[Any] = []
        for item in primary_constraints:
            if isinstance(item, dict):
                enriched = dict(item)
                enriched.setdefault("constraint_source", "primary")
                merged_constraints.append(enriched)
        for item in harvested_constraints:
            if isinstance(item, dict):
                enriched = dict(item)
                enriched.setdefault("constraint_source", "harvested")
                merged_constraints.append(enriched)

        constraints = merged_constraints if merged_constraints else constraints_raw

        # Normalise constraint keys for downstream compatibility
        if constraints_is_string:
            # Preserve the expression verbatim for downstream export/validation
            pass
        else:
            normalised_constraints: List[Dict[str, Any]] = []
            for c in constraints:
                if not isinstance(c, dict):
                    continue
                d = dict(c)
                if "subject" not in d and isinstance(d.get("parameter"), (str, int, float)):
                    d["subject"] = d.get("parameter")
                if "unit" not in d and isinstance(d.get("units"), (str, int, float)):
                    d["unit"] = d.get("units")
                if "logic" not in d and d.get("expression") is not None:
                    d["logic"] = d.get("expression")
                normalised_constraints.append(d)
            constraints = normalised_constraints

        raw_severity = str(payload.get("severity") or payload.get("rule_severity") or "").strip().upper()
        severity = raw_severity if raw_severity in {"ENFORCEABLE", "ADVISORY"} else ""

        has_constraints = bool(constraints) and (constraints.strip() != "" if isinstance(constraints, str) else True)
        compound_logic = payload.get("compound_logic")
        if not isinstance(compound_logic, list):
            compound_logic = []

        suggested_severity = "ENFORCEABLE" if has_constraints else "ADVISORY"

        supporting_quote = str(payload.get("supporting_quote", "")).strip()
        supporting_quotes = payload.get("supporting_quotes") or payload.get("supportingQuotes")
        if isinstance(supporting_quotes, list):
            supporting_quotes_list = [str(item).strip() for item in supporting_quotes if str(item).strip()]
        else:
            supporting_quotes_list = []
        if supporting_quote and supporting_quote not in supporting_quotes_list:
            supporting_quotes_list.insert(0, supporting_quote)

        metadata = {
            "source_document": document_name,
            "chunk_index": chunk_index,
            "chunk_rule_index": rule_index,
            "canonical_statement": canonical or raw_text,
            "domain": domain_text.strip() or None,
            "primary_feature": str(payload.get("primary_feature", "")).strip(),
            "unit": str(payload.get("unit", "")).strip(),
            "value": payload.get("value"),
            "tolerance_range": payload.get("tolerance_range"),
            "supporting_quote": supporting_quote,
            "intent": intent,
            "scope_domain": scope_domain,
            "applicability": applicability,
            "applicability_structured": app_struct if isinstance(app_struct, dict) else None,
            "constraints": constraints,
            "has_constraints": has_constraints,
            "severity": severity or suggested_severity,
            "suggested_severity": suggested_severity,
            "compound_logic": compound_logic,
            "supporting_quotes": supporting_quotes_list,
            "provenance": {
                "source_document": document_name,
                "chunk_index": chunk_index,
                "chunk_rule_index": rule_index,
                "supporting_quotes": supporting_quotes_list,
                "section_title": payload.get("section_title") or payload.get("section") or None,
            },
        }

        rationale = str(payload.get("rationale", "")).strip()
        return Rule(
            rule_text=raw_text,
            category=category,
            rule_type=rule_type,
            confidence=confidence,
            priority=priority,
            rationale=rationale,
            metadata=metadata,
        )


# ---------------------------------------------------------------------------
# Public pipeline
# ---------------------------------------------------------------------------


class RuleExtractionPipeline:
    """Coordinates document loading, chunking, and Groq calls."""

    def __init__(
        self,
        settings: Optional[RuleExtractionSettings] = None,
        *,
        loader: Optional[DocumentLoader] = None,
        chunker: Optional[TextChunker] = None,
    ) -> None:
        self.settings = settings or RuleExtractionSettings()
        self.loader = loader or DocumentLoader(cache_size=self.settings.document_cache_size)
        self.chunker = chunker or TextChunker()
        self.processor = ChunkProcessor(self.settings)

    async def extract_from_text(
        self,
        text: str,
        document_name: str,
        *,
        max_rules: Optional[int] = None,
    ) -> ExtractionSummary:
        start = time.perf_counter()
        cleaned_text = (text or "").strip()

        if len(cleaned_text) < self.settings.min_text_length:
            processing_time = time.perf_counter() - start
            return ExtractionSummary(
                filename=document_name,
                status="insufficient_text",
                rules=[],
                rule_count=0,
                chunks_processed=0,
                model_calls=0,
                processing_time=processing_time,
                avg_confidence=0.0,
            )

        chunks = self._split_text_fast(cleaned_text)
        limit = max_rules or self.settings.max_rules_total
        aggregated: "OrderedDict[str, Rule]" = OrderedDict()
        model_calls = 0

        tasks: List[asyncio.Task[List[Rule]]] = []
        semaphore = asyncio.Semaphore(self.settings.max_concurrent_calls)
        cancelled = False

        async def worker(chunk_index: int, chunk_text: str) -> List[Rule]:
            nonlocal model_calls
            async with semaphore:
                if cancelled:
                    return []
                model_calls += 1
                return await self.processor.process(
                    chunk_text,
                    document_name=document_name,
                    chunk_index=chunk_index,
                )

        for index, chunk in enumerate(chunks):
            tasks.append(asyncio.create_task(worker(index, chunk)))

        try:
            for task in asyncio.as_completed(tasks):
                rules = await task
                for rule in rules:
                    key = self._dedupe_key(rule)
                    if key not in aggregated:
                        aggregated[key] = rule
                    if len(aggregated) >= limit:
                        cancelled = True
                        for pending in tasks:
                            if not pending.done():
                                pending.cancel()
                        break
                if cancelled:
                    break
        finally:
            for task in tasks:
                if not task.done():
                    task.cancel()

        final_rules = list(aggregated.values())[:limit]
        avg_confidence = (
            sum(rule.confidence for rule in final_rules) / len(final_rules)
            if final_rules
            else 0.0
        )
        processing_time = time.perf_counter() - start

        return ExtractionSummary(
            filename=document_name,
            status="success" if final_rules else "no_rules",
            rules=[rule.to_dict() for rule in final_rules],
            rule_count=len(final_rules),
            chunks_processed=len(chunks),
            model_calls=model_calls,
            processing_time=processing_time,
            avg_confidence=round(avg_confidence, 4),
        )

    async def process_document(
        self,
        document_path: str,
        *,
        max_rules: Optional[int] = None,
    ) -> ExtractionSummary:
        path = Path(document_path)
        payload = await self.loader.load(path)
        summary = await self.extract_from_text(
            payload.text,
            payload.metadata.get("filename", path.name),
            max_rules=max_rules,
        )
        return replace(summary, document_metadata=payload.metadata)

    async def batch_process(
        self,
        document_paths: Sequence[str],
        *,
        max_rules: Optional[int] = None,
        concurrency: Optional[int] = None,
    ) -> List[ExtractionSummary]:
        if not document_paths:
            return []

        limit = concurrency or self.settings.max_documents_concurrency
        semaphore = asyncio.Semaphore(max(1, limit))
        summaries: List[ExtractionSummary] = []

        async def worker(path: str) -> None:
            async with semaphore:
                result = await self.process_document(path, max_rules=max_rules)
                summaries.append(result)

        await asyncio.gather(*(worker(path) for path in document_paths))
        return summaries

    # ------------------------------ internals ------------------------------

    def _split_text_fast(self, text: str) -> List[str]:
        try:
            return self.chunker.split(
                text,
                window=self.settings.max_chunk_tokens,
                overlap=self.settings.chunk_overlap_tokens,
            )
        except Exception:  # pragma: no cover - tiktoken edge cases
            return self.chunker.word_fallback(
                text,
                window=self.settings.max_chunk_tokens,
                overlap=self.settings.chunk_overlap_tokens,
            )

    def _dedupe_key(self, rule: Rule) -> str:
        def _norm_text(value: Any) -> str:
            return re.sub(r"\s+", " ", str(value or "").strip().lower())

        intent = _norm_text(rule.metadata.get("intent"))
        scope_domain = rule.metadata.get("scope_domain")
        if isinstance(scope_domain, dict):
            scope_sig = "|".join(
                _norm_text(scope_domain.get(key))
                for key in ("process", "item", "feature")
            )
        else:
            scope_sig = ""

        applicability = rule.metadata.get("applicability")
        if isinstance(applicability, list):
            app_sig = ";".join(
                _norm_text(
                    (
                        item.get("exp_name")
                        or item.get("gate")
                        or item.get("name")
                        or item.get("field")
                        or item.get("attribute")
                        or item.get("subject")
                    )
                    if isinstance(item, dict)
                    else item
                )
                + ":"
                + _norm_text(
                    (item.get("operator") or item.get("op"))
                    if isinstance(item, dict)
                    else ""
                )
                + ":"
                + _norm_text(
                    (
                        item.get("value")
                        if isinstance(item, dict)
                        else ""
                    )
                )
                for item in applicability
                if item
            )
        else:
            app_sig = ""

        constraints = rule.metadata.get("constraints")
        if isinstance(constraints, list):
            con_sig = ";".join(
                _norm_text(item.get("subject") if isinstance(item, dict) else item)
                + ":"
                + _norm_text(item.get("operator") if isinstance(item, dict) else "")
                + ":"
                + _norm_text(item.get("value") if isinstance(item, dict) else "")
                + ":"
                + _norm_text(item.get("unit") if isinstance(item, dict) else "")
                + ":"
                + _norm_text(item.get("logic") if isinstance(item, dict) else "")
                for item in constraints
                if item
            )
        else:
            con_sig = ""

        if intent or app_sig or con_sig or scope_sig:
            return f"{intent}|{scope_sig}|{app_sig}|{con_sig}"

        base = _norm_text(rule.rule_text)
        return f"{base}|{_norm_text(rule.category)}|{_norm_text(rule.rule_type)}"


# ---------------------------------------------------------------------------
# Convenience helpers used by legacy wrappers
# ---------------------------------------------------------------------------


async def extract_rules_from_text(
    text: str,
    document_name: str,
    *,
    settings: Optional[RuleExtractionSettings] = None,
    max_rules: Optional[int] = None,
) -> ExtractionSummary:
    pipeline = RuleExtractionPipeline(settings)
    return await pipeline.extract_from_text(text, document_name, max_rules=max_rules)


async def process_document(
    document_path: str,
    *,
    settings: Optional[RuleExtractionSettings] = None,
    max_rules: Optional[int] = None,
) -> ExtractionSummary:
    pipeline = RuleExtractionPipeline(settings)
    return await pipeline.process_document(document_path, max_rules=max_rules)


async def batch_process(
    document_paths: Sequence[str],
    *,
    settings: Optional[RuleExtractionSettings] = None,
    max_rules: Optional[int] = None,
    concurrency: Optional[int] = None,
) -> List[ExtractionSummary]:
    pipeline = RuleExtractionPipeline(settings)
    return await pipeline.batch_process(
        document_paths,
        max_rules=max_rules,
        concurrency=concurrency,
    )
